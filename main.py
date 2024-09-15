import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from docx import Document
from io import BytesIO
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv

load_dotenv()
GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
genai.configure(api_key=GOOGLE_API_KEY)


# Initialize Vertex AI model
#
# class DocumentProcessor:
#     @staticmethod
#     def extract_text_from_pdf(file):
#         text = ""
#         pdf_reader = PdfReader(BytesIO(file.read()))
#         for page in pdf_reader.pages:
#             text += page.extract_text()
#         return text
#
#     @staticmethod
#     def extract_text_from_docx(file):
#         doc = Document(BytesIO(file.read()))
#         full_text = [para.text for para in doc.paragraphs]
#         return '\n'.join(full_text)
#
#     @staticmethod
#     def get_text(files):
#         if not files:
#             return False
#
#         uploaded_file = files[0]
#
#         if "pdf" in uploaded_file.name.lower():
#             return DocumentProcessor.extract_text_from_pdf(uploaded_file)
#         elif any(ext in uploaded_file.name.lower() for ext in ("docx", "doc")):
#             return DocumentProcessor.extract_text_from_docx(uploaded_file)
#         else:
#             return False
#
#
class VectorStoreManager:
    #     @staticmethod
    #     def get_text_chunks(text):
    #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    #         return text_splitter.split_text(text)
    #
    #     @staticmethod
    #     def create_vector_store(text_chunks):
    #         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    #         vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    #         vector_store.save_local("faiss_index")
    #
    @staticmethod
    def load_vector_store():
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        index_path = "faiss_index"
        if os.path.exists(index_path):
            return FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        else:
            return None


class ChatManager:
    @staticmethod
    def get_conversational_chain():
        prompt_template = """
        Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
        provided context just say "answer is not available in the context" and don't provide the wrong answer\n\n
        Context:\n {context}?\n
        Question: \n{question}\n

        Answer:
        """

        model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
        prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
        return load_qa_chain(model, chain_type="stuff", prompt=prompt)

    @staticmethod
    def get_response(user_question):
        if vector_store := VectorStoreManager.load_vector_store():

            docs = vector_store.similarity_search(user_question, k=10)

            chain = ChatManager.get_conversational_chain()

            response = chain(
                {"input_documents": docs, "question": user_question},
                return_only_outputs=True
            )

            if "answer is not available in the context" in response["output_text"].lower():
                st.write("Context-specific answer not found. Fetching a general answer.")
                general_answer = genai.GenerativeModel("gemini-1.5-flash").generate_content(user_question)
                return general_answer.text
            return response["output_text"]
        else:
            st.write("No Docs are available.Data is fetching from general context...")
            general_answer = genai.GenerativeModel("gemini-1.5-flash").generate_content(user_question)
            return general_answer.text


def main():
    st.set_page_config(page_title="Chat with Documents", page_icon="💁", layout="wide")

    st.title("Chat with Your Documents 💁")
    st.markdown("""

        <style>
        /* Hide the GitHub icon in the Streamlit app */
           #MainMenu {visibility: hidden;}
            footer {visibility:hidden;}
            .GithubIcon {visibility: hidden;}
        .intro-text {
            font-size: 18px;
            color: #4F4F4F;
            text-align: center;
            margin: 20px 0;
        }
        .highlight {
            font-weight: bold;
            color: #FF6347;
        }
        </style>
        <div class="intro-text">
            Upload your <span class="highlight">PDF</span> or <span class="highlight">DOCX</span> files,
            and then ask any questions you have about the content.
            The system will provide detailed answers based on the uploaded documents.
        </div>
        """, unsafe_allow_html=True)

    # with st.sidebar:
    # st.title("Upload and Process Documents")
    # pdf_docs = st.file_uploader("Upload your PDF or DOCX files", accept_multiple_files=True)

    # if st.button("Submit & Process"):
    #     with st.spinner("Processing..."):
    #         raw_text = DocumentProcessor.get_text(pdf_docs)
    #         if raw_text:
    #             text_chunks = VectorStoreManager.get_text_chunks(raw_text)
    #             VectorStoreManager.create_vector_store(text_chunks)
    #             st.success("Documents processed successfully!")
    #         else:
    #             st.error("Please upload a valid PDF or DOCX file.")

    st.subheader("Ask a Question")
    user_question = st.text_input("Type your question here:")

    if user_question:
        with st.spinner("Generating response..."):
            response = ChatManager.get_response(user_question)
            st.markdown("### Reply:")
            st.markdown(response)


if __name__ == "__main__":
    main()

#
# # part2_question_answering.py
#
# import streamlit as st
# import faiss
# import numpy as np
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.chains.question_answering import load_qa_chain
# from langchain.prompts import PromptTemplate
# from dotenv import load_dotenv
# from langchain_google_genai import GoogleGenerativeAIEmbeddings
#
# load_dotenv()
# GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
#
# class ChatManager:
#     @staticmethod
#     def load_faiss_index(index_path):
#         # Load FAISS index from file
#         return faiss.read_index(index_path)
#
#     @staticmethod
#     def get_conversational_chain():
#         prompt_template = """
#         Answer the question as detailed as possible from the provided context. If the answer is not in the provided context, just say "answer is not available in the context" and don't provide a wrong answer.\n\n
#         Context:\n{context}\n
#         Question: \n{question}\n
#
#         Answer:
#         """
#
#         model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
#         prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
#         return load_qa_chain(model, chain_type="stuff", prompt=prompt)
#
#     @staticmethod
#     def get_response(user_question):
#         # Path to the FAISS index file
#         index_path = 'faiss_index'
#         index = ChatManager.load_faiss_index(index_path)
#
#         # Assuming you have the embeddings model to convert query into vector
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
#         query_vector = embeddings.embed([user_question])[0]
#         query_vector = np.array(query_vector).astype(np.float32).reshape(1, -1)
#
#         # Search the FAISS index
#         D, I = index.search(query_vector, k=10)
#
#         # Extract top results
#         # Replace this with your own logic to retrieve documents based on indices
#         docs = ["Document " + str(i) for i in I[0]]  # Dummy document retrieval
#
#         chain = ChatManager.get_conversational_chain()
#         response = chain(
#             {"input_documents": docs, "question": user_question},
#             return_only_outputs=True
#         )
#
#         if "answer is not available in the context" in response["output_text"].lower():
#             st.write("Context-specific answer not found. Fetching a general answer.")
#             general_answer = genai.GenerativeModel("gemini-1.5-flash").generate_content(user_question)
#             return general_answer.text
#         return response["output_text"]
#
# def main():
#     st.set_page_config(page_title="Ask Questions", page_icon="❓")
#
#     st.title("Ask Questions Based on Indexed Documents")
#     st.markdown("Type your question below to get answers based on the indexed documents.")
#
#     st.subheader("Ask a Question")
#     user_question = st.text_input("Type your question here:")
#
#     if user_question:
#         with st.spinner("Generating response..."):
#             response = ChatManager.get_response(user_question)
#             st.markdown("### Reply:")
#             st.markdown(response)
#
# if __name__ == "__main__":
#     main()

#
# # part1_index_creation.py
#
# import streamlit as st
# from PyPDF2 import PdfReader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import os
# from langchain_google_genai import GoogleGenerativeAIEmbeddings
# from langchain_community.vectorstores import FAISS
# from docx import Document
# from io import BytesIO
# from dotenv import load_dotenv
#
# load_dotenv()
# GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
#
# class DocumentProcessor:
#     @staticmethod
#     def extract_text_from_pdf(file):
#         text = ""
#         pdf_reader = PdfReader(BytesIO(file.read()))
#         for page in pdf_reader.pages:
#             text += page.extract_text()
#         return text
#
#     @staticmethod
#     def extract_text_from_docx(file):
#         doc = Document(BytesIO(file.read()))
#         full_text = [para.text for para in doc.paragraphs]
#         return '\n'.join(full_text)
#
#     @staticmethod
#     def get_text(files):
#         text = ""
#         for uploaded_file in files:
#             if "pdf" in uploaded_file.name.lower():
#                 text += DocumentProcessor.extract_text_from_pdf(uploaded_file)
#             elif any(ext in uploaded_file.name.lower() for ext in ("docx", "doc")):
#                 text += DocumentProcessor.extract_text_from_docx(uploaded_file)
#         return text if text else False
#
# class VectorStoreManager:
#     @staticmethod
#     def get_text_chunks(text):
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
#         return text_splitter.split_text(text)
#
#     @staticmethod
#     def create_vector_store(text_chunks):
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
#         vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
#         vector_store.save_local("faiss_index")
#
# def main():
#     st.set_page_config(page_title="Create FAISS Index", page_icon="🗂️")
#
#     st.title("Create FAISS Index from Documents")
#     st.markdown("Upload your PDF or DOCX files to create the FAISS index.")
#
#     with st.sidebar:
#         st.title("Upload and Process Documents")
#         pdf_docs = st.file_uploader("Upload your PDF or DOCX files", accept_multiple_files=True)
#
#         if st.button("Submit & Process"):
#             with st.spinner("Processing..."):
#                 raw_text = DocumentProcessor.get_text(pdf_docs)
#                 if raw_text:
#                     text_chunks = VectorStoreManager.get_text_chunks(raw_text)
#                     VectorStoreManager.create_vector_store(text_chunks)
#                     st.success("Documents processed successfully and FAISS index created!")
#                 else:
#                     st.error("Please upload valid PDF or DOCX files.")
#
# if __name__ == "__main__":
#     main()
